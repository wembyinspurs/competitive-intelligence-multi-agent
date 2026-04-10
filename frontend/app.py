"""竞品情报分析系统可视化前端"""
import streamlit as st
import requests
import json
from sseclient import SSEClient
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io

# 页面配置
st.set_page_config(
    page_title="多Agent竞品情报分析系统",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 全局配置
API_BASE_URL = "http://localhost:8000"

# 侧边栏导航
with st.sidebar:
    st.title("📊 竞品情报系统")
    st.divider()
    page = st.radio("功能导航", ["竞品分析工作台", "竞品管理", "历史分析", "系统配置"])
    st.divider()
    st.caption("基于多Agent架构的企业级竞品情报分析平台")

# 健康检查
def check_health():
    try:
        res = requests.get(f"{API_BASE_URL}/health", timeout=5)
        return res.status_code == 200
    except:
        return False

# 获取竞品列表
@st.cache_data(ttl=60)
def get_competitor_list():
    try:
        res = requests.get(f"{API_BASE_URL}/competitors/all", timeout=5)
        if res.status_code == 200:
            return res.json()
        return []
    except:
        return []

# 生成Word竞品分析报告
def generate_word_report(competitor_name, analysis_result, quality_score):
    doc = Document()
    # 标题
    doc.add_heading(f"{competitor_name} 竞品分析报告", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"分析质量评分：{quality_score}/10")
    doc.add_page_break()

    # 1. 竞品对比矩阵
    doc.add_heading("一、竞品对比矩阵", level=1)
    if analysis_result.get("comparison_matrix"):
        matrix = analysis_result["comparison_matrix"]
        for dim in matrix["dimensions"]:
            doc.add_heading(dim["dimension"], level=2)
            doc.add_paragraph(f"我方评分：{dim['our_score']}/10 | 竞品评分：{dim['competitor_score']}/10")
            doc.add_paragraph(f"分析说明：{dim['notes']}")
        doc.add_heading("整体评估", level=2)
        doc.add_paragraph(matrix["overall_assessment"])
    doc.add_page_break()

    # 2. 销售战术卡
    doc.add_heading("二、销售战术卡", level=1)
    if analysis_result.get("battlecard"):
        card = analysis_result["battlecard"]
        doc.add_heading("我方优势", level=2)
        for item in card["our_strengths"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("我方劣势", level=2)
        for item in card["our_weaknesses"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("竞品优势", level=2)
        for item in card["competitor_strengths"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("竞品劣势", level=2)
        for item in card["competitor_weaknesses"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("核心差异化", level=2)
        for item in card["key_differentiators"]:
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("异议处理话术", level=2)
        for question, answer in card["objection_handling"].items():
            doc.add_paragraph(f"Q: {question}", style="List Bullet")
            doc.add_paragraph(f"A: {answer}")
        doc.add_heading("电梯Pitch", level=2)
        doc.add_paragraph(card["elevator_pitch"])
    doc.add_page_break()

    # 3. 研究洞察
    doc.add_heading("三、研究洞察", level=1)
    if analysis_result.get("research_results"):
        for insight in analysis_result["research_results"]:
            doc.add_heading(insight["topic"], level=2)
            doc.add_paragraph(insight["summary"])
            doc.add_heading("核心发现", level=3)
            for finding in insight["key_findings"]:
                doc.add_paragraph(f"- {finding}", style="List Bullet")
    doc.add_page_break()

    # 4. 变更监控
    doc.add_heading("四、变更监控", level=1)
    if analysis_result.get("changes_detected"):
        for change in analysis_result["changes_detected"]:
            doc.add_paragraph(f"[{change['severity']}] {change['title']}")
            doc.add_paragraph(change["summary"])
    else:
        doc.add_paragraph("本次监控未检测到竞品变更")

    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ------------------------------
# 页面1：竞品分析工作台（核心）
# ------------------------------
if page == "竞品分析工作台":
    st.header("🔍 竞品全维度分析工作台")
    st.caption("输入竞品信息，一键启动多Agent智能分析")

    # 服务状态检查
    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动，请先运行 uvicorn src.api.server:app --reload --port 8000")
    else:
        st.success("✅ 后端服务正常运行")

    # 新增：快速选择已有竞品
    competitor_list = get_competitor_list()
    use_exist_competitor = st.checkbox("选择已有竞品", value=False)
    selected_competitor = None
    competitor_name = ""
    competitor_urls = ""

    if use_exist_competitor and competitor_list:
        competitor_options = {comp["name"]: comp for comp in competitor_list}
        selected_name = st.selectbox("选择竞品", options=list(competitor_options.keys()))
        selected_competitor = competitor_options[selected_name]
        competitor_name = selected_competitor["name"]
        competitor_urls = "\n".join(selected_competitor["urls"])
        st.info(f"已选择竞品：{selected_name}，自动填充监控URL")

    # 输入区域
    with st.form("analyze_form"):
        col1, col2 = st.columns([1, 2])
        with col1:
            competitor_name = st.text_input("竞品名称", value=competitor_name, placeholder="例如：字节跳动")
        with col2:
            competitor_urls = st.text_area(
                "监控URL（一行一个）",
                value=competitor_urls,
                placeholder="https://www.bytedance.com\nhttps://www.feishu.cn",
                height=80
            )
        submit = st.form_submit_button("🚀 启动全流程分析", use_container_width=True, type="primary")

    # 分析执行逻辑
    if submit and service_online:
        if not competitor_name:
            st.warning("请输入竞品名称")
        else:
            # 处理URL
            url_list = [u.strip() for u in competitor_urls.split("\n") if u.strip()]
            request_body = {
                "competitor": competitor_name,
                "urls": url_list
            }

            # 进度展示区域
            st.divider()
            st.subheader("📡 实时执行进度")
            progress_container = st.container()
            result_container = st.container()

            # 调用SSE流式接口
            try:
                response = requests.post(
                    f"{API_BASE_URL}/analyze/stream",
                    json=request_body,
                    stream=True,
                    headers={"Accept": "text/event-stream"}
                )
                client = SSEClient(response)

                # 存储各节点结果
                node_results = {}
                final_analysis_result = {}
                final_quality_score = 0.0

                # 实时处理事件
                for event in client.events():
                    if event.event == "ping":
                        continue
                    if event.event == "error":
                        progress_container.error(f"❌ 执行出错：{json.loads(event.data)['error']}")
                        break

                    # 展示节点执行进度
                    node_name = event.event
                    node_data = json.loads(event.data)
                    node_results[node_name] = node_data

                    with progress_container:
                        with st.expander(f"✅ {node_name} 节点执行完成", expanded=True):
                            st.json(node_data, expanded=False)

                    # 保存最终结果
                    if node_name == "quality_check":
                        final_quality_score = node_data.get("quality_score", 0.0)
                    if node_name == "battlecard":
                        final_analysis_result["battlecard"] = node_data.get("battlecard")
                    if node_name == "compare":
                        final_analysis_result["comparison_matrix"] = node_data.get("comparison_matrix")
                    if node_name == "research":
                        final_analysis_result["research_results"] = node_data.get("research_results", [])
                    if node_name == "monitor":
                        final_analysis_result["changes_detected"] = node_data.get("changes_detected", [])

                # 分析完成，展示最终结果
                with result_container:
                    st.divider()
                    st.header("📑 最终分析报告")

                    # 新增：Word报告导出按钮
                    col_export, col_score = st.columns([1, 1])
                    with col_export:
                        if final_analysis_result:
                            word_buffer = generate_word_report(competitor_name, final_analysis_result, final_quality_score)
                            st.download_button(
                                label="📥 导出Word分析报告",
                                data=word_buffer,
                                file_name=f"{competitor_name}_竞品分析报告_{datetime.now().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    with col_score:
                        st.metric("最终质量评分", f"{final_quality_score}/10")

                    # 分Tab展示结果
                    tab1, tab2, tab3, tab4 = st.tabs([
                        "📈 对比矩阵", "🎯 销售战术卡", "🔎 研究洞察", "⚠️ 变更监控"
                    ])

                    # 对比矩阵
                    with tab1:
                        if "comparison_matrix" in final_analysis_result and final_analysis_result["comparison_matrix"]:
                            matrix = final_analysis_result["comparison_matrix"]
                            st.subheader(f"竞品：{matrix['competitor']}")
                            for dim in matrix["dimensions"]:
                                col1, col2, col3 = st.columns([2, 1, 1])
                                with col1:
                                    st.markdown(f"**{dim['dimension']}**")
                                    st.caption(dim["notes"])
                                with col2:
                                    st.metric("我方评分", dim["our_score"])
                                with col3:
                                    st.metric("竞品评分", dim["competitor_score"])
                                st.divider()
                            st.markdown(f"**整体评估**：{matrix['overall_assessment']}")

                    # 销售战术卡
                    with tab2:
                        if "battlecard" in final_analysis_result and final_analysis_result["battlecard"]:
                            card = final_analysis_result["battlecard"]
                            col1, col2 = st.columns(2)
                            with col1:
                                st.subheader("✅ 我方优势")
                                for item in card["our_strengths"]:
                                    st.write(f"- {item}")
                                st.subheader("⚠️ 我方劣势")
                                for item in card["our_weaknesses"]:
                                    st.write(f"- {item}")
                            with col2:
                                st.subheader("📈 竞品优势")
                                for item in card["competitor_strengths"]:
                                    st.write(f"- {item}")
                                st.subheader("❌ 竞品劣势")
                                for item in card["competitor_weaknesses"]:
                                    st.write(f"- {item}")
                            st.divider()
                            st.subheader("🎯 核心差异化")
                            for item in card["key_differentiators"]:
                                st.write(f"- {item}")
                            st.divider()
                            st.subheader("🗣️ 异议处理话术")
                            for question, answer in card["objection_handling"].items():
                                with st.expander(question):
                                    st.write(answer)
                            st.divider()
                            st.subheader("🎙️ 电梯Pitch")
                            st.info(card["elevator_pitch"])

                    # 研究洞察
                    with tab3:
                        if "research_results" in final_analysis_result:
                            insights = final_analysis_result["research_results"]
                            for insight in insights:
                                with st.expander(f"📝 {insight['topic']}", expanded=True):
                                    st.write(insight["summary"])
                                    st.markdown("**核心发现**")
                                    for finding in insight["key_findings"]:
                                        st.write(f"- {finding}")

                    # 变更监控
                    with tab4:
                        if "changes_detected" in final_analysis_result:
                            changes = final_analysis_result["changes_detected"]
                            if not changes:
                                st.info("本次监控未检测到竞品变更")
                            else:
                                for change in changes:
                                    st.write(f"[{change['severity']}] {change['title']}")

            except Exception as e:
                st.error(f"请求失败：{str(e)}")

# ------------------------------
# 页面2：竞品管理
# ------------------------------
elif page == "竞品管理":
    st.header("📋 竞品库管理")
    st.caption("统一管理竞品信息，支持新增、编辑、删除、快速选择")

    # 服务状态检查
    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动，请先运行 uvicorn src.api.server:app --reload --port 8000")
    else:
        # 刷新竞品列表
        if st.button("🔄 刷新列表", use_container_width=False):
            st.cache_data.clear()
            st.rerun()

        # 新增竞品表单
        with st.expander("➕ 新增竞品", expanded=False):
            with st.form("add_competitor_form"):
                col1, col2 = st.columns([1, 2])
                with col1:
                    new_name = st.text_input("竞品名称", placeholder="例如：字节跳动")
                with col2:
                    new_urls = st.text_area(
                        "监控URL（一行一个）",
                        placeholder="https://www.bytedance.com\nhttps://www.feishu.cn",
                        height=80
                    )
                add_submit = st.form_submit_button("确认新增", use_container_width=True, type="primary")

                if add_submit:
                    if not new_name:
                        st.warning("请输入竞品名称")
                    else:
                        url_list = [u.strip() for u in new_urls.split("\n") if u.strip()]
                        res = requests.post(
                            f"{API_BASE_URL}/competitors",
                            json={"name": new_name, "urls": url_list}
                        )
                        if res.status_code == 200:
                            st.success(f"✅ 竞品【{new_name}】新增成功")
                            st.cache_data.clear()
                            st.rerun()
                        else:
                            st.error(f"❌ 新增失败：{res.json()['detail']}")

        # 竞品列表展示
        st.subheader("竞品列表")
        competitor_list = get_competitor_list()
        if not competitor_list:
            st.info("暂无竞品数据，点击上方「新增竞品」添加")
        else:
            for comp in competitor_list:
                with st.container(border=True):
                    col1, col2, col3 = st.columns([2, 3, 1])
                    with col1:
                        st.markdown(f"### {comp['name']}")
                        st.caption(f"创建时间：{comp['created_at'].split('T')[0]}")
                    with col2:
                        st.markdown("**监控URL**")
                        for url in comp["urls"]:
                            st.code(url, language="text")
                    with col3:
                        # 编辑按钮
                        if st.button("✏️ 编辑", key=f"edit_{comp['id']}", use_container_width=True):
                            st.session_state[f"edit_comp_{comp['id']}"] = True
                        # 删除按钮
                        if st.button("🗑️ 删除", key=f"del_{comp['id']}", use_container_width=True, type="secondary"):
                            res = requests.delete(f"{API_BASE_URL}/competitors/{comp['id']}")
                            if res.status_code == 200:
                                st.success("✅ 删除成功")
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error(f"❌ 删除失败：{res.json()['detail']}")

                    # 编辑表单
                    if st.session_state.get(f"edit_comp_{comp['id']}", False):
                        with st.form(f"edit_form_{comp['id']}"):
                            edit_name = st.text_input("竞品名称", value=comp["name"])
                            edit_urls = st.text_area(
                                "监控URL（一行一个）",
                                value="\n".join(comp["urls"]),
                                height=80
                            )
                            col_cancel, col_confirm = st.columns(2)
                            with col_cancel:
                                if st.form_submit_button("取消", use_container_width=True):
                                    st.session_state[f"edit_comp_{comp['id']}"] = False
                                    st.rerun()
                            with col_confirm:
                                edit_submit = st.form_submit_button("确认修改", use_container_width=True, type="primary")
                                if edit_submit:
                                    if not edit_name:
                                        st.warning("请输入竞品名称")
                                    else:
                                        url_list = [u.strip() for u in edit_urls.split("\n") if u.strip()]
                                        res = requests.put(
                                            f"{API_BASE_URL}/competitors/{comp['id']}",
                                            json={"name": edit_name, "urls": url_list}
                                        )
                                        if res.status_code == 200:
                                            st.success("✅ 修改成功")
                                            st.session_state[f"edit_comp_{comp['id']}"] = False
                                            st.cache_data.clear()
                                            st.rerun()
                                        else:
                                            st.error(f"❌ 修改失败：{res.json()['detail']}")

# ------------------------------
# 页面3：历史分析
# ------------------------------
elif page == "历史分析":
    st.header("📂 历史分析回溯")
    st.caption("查看历史分析报告，对比竞品变化趋势")

    # 服务状态检查
    service_online = check_health()
    if not service_online:
        st.error("❌ 后端服务未启动，请先运行 uvicorn src.api.server:app --reload --port 8000")
    else:
        # 筛选条件
        competitor_list = get_competitor_list()
        competitor_options = {"全部竞品": None}
        for comp in competitor_list:
            competitor_options[comp["name"]] = comp["id"]
        selected_competitor = st.selectbox("按竞品筛选", options=list(competitor_options.keys()))
        selected_competitor_id = competitor_options[selected_competitor]

        # 获取分析记录
        res = requests.get(
            f"{API_BASE_URL}/analysis/records",
            params={"competitor_id": selected_competitor_id} if selected_competitor_id else {}
        )
        if res.status_code != 200:
            st.error("❌ 获取历史记录失败")
            record_list = []
        else:
            record_list = res.json()

        # 记录列表展示
        if not record_list:
            st.info("暂无历史分析记录，请到「竞品分析工作台」执行分析")
        else:
            for record in record_list:
                with st.container(border=True):
                    col1, col2, col3 = st.columns([2, 1, 1])
                    with col1:
                        st.markdown(f"### 竞品：{record['competitor_name']}")
                        st.caption(f"分析时间：{record['created_at'].replace('T', ' ').split('.')[0]}")
                    with col2:
                        st.metric("质量评分", f"{record['quality_score']}/10")
                    with col3:
                        if st.button("📄 查看详情", key=f"view_{record['id']}", use_container_width=True, type="primary"):
                            st.session_state[f"view_record_{record['id']}"] = True

                    # 详情展示
                    if st.session_state.get(f"view_record_{record['id']}", False):
                        with st.expander("分析详情", expanded=True):
                            result = record["analysis_result"]
                            tab1, tab2, tab3, tab4 = st.tabs([
                                "📈 对比矩阵", "🎯 销售战术卡", "🔎 研究洞察", "⚠️ 变更监控"
                            ])
                            with tab1:
                                if result.get("comparison_matrix"):
                                    matrix = result["comparison_matrix"]
                                    for dim in matrix["dimensions"]:
                                        st.markdown(f"**{dim['dimension']}** | 我方：{dim['our_score']}/10 | 竞品：{dim['competitor_score']}/10")
                                        st.caption(dim["notes"])
                            with tab2:
                                if result.get("battlecard"):
                                    card = result["battlecard"]
                                    st.markdown("**核心差异化**")
                                    for item in card["key_differentiators"]:
                                        st.write(f"- {item}")
                            with tab3:
                                if result.get("research_results"):
                                    for insight in result["research_results"]:
                                        st.markdown(f"**{insight['topic']}**")
                                        st.write(insight["summary"])
                            with tab4:
                                if result.get("changes_detected"):
                                    for change in result["changes_detected"]:
                                        st.write(f"[{change['severity']}] {change['title']}")

                            # 关闭详情按钮
                            if st.button("关闭详情", key=f"close_{record['id']}", use_container_width=False):
                                st.session_state[f"view_record_{record['id']}"] = False
                                st.rerun()

# ------------------------------
# 页面4：系统配置
# ------------------------------
elif page == "系统配置":
    st.header("⚙️ 系统配置")
    st.info("功能开发中：支持告警阈值、通知渠道、LLM参数配置")